# --- Setup and imports ---

import os
from pathlib import Path
import json
import re
import uuid
from typing import List, Dict, TypedDict, Optional

import pandas as pd
from docx import Document
from dotenv import load_dotenv

from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.schema import Document as LCDocument
from langgraph.graph import StateGraph, START, END
from tavily import TavilyClient
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import markdown as md
from weasyprint import HTML, CSS



load_dotenv()

DATA_DIR = Path("data")
REPORTS_DIR = Path("reports"); REPORTS_DIR.mkdir(exist_ok=True)
VSTORE_DIR = Path("vectorstores"); VSTORE_DIR.mkdir(exist_ok=True)

AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
DEPLOYMENT_LLM = os.getenv("AZURE_OPENAI_DEPLOYMENT_LLM")
DEPLOYMENT_EMB = os.getenv("AZURE_OPENAI_DEPLOYMENT_EMBEDDING")
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

# LLMs
llm = AzureChatOpenAI(
    azure_deployment=DEPLOYMENT_LLM,
    openai_api_key=AZURE_API_KEY,
    azure_endpoint=AZURE_ENDPOINT,
    openai_api_version=AZURE_API_VERSION,
    model_kwargs={"max_completion_tokens": 3000}
)

embeddings = AzureOpenAIEmbeddings(
    azure_deployment=DEPLOYMENT_EMB,
    openai_api_key=AZURE_API_KEY,
    azure_endpoint=AZURE_ENDPOINT,
    openai_api_version=AZURE_API_VERSION,
)

tavily_client = TavilyClient(api_key=TAVILY_API_KEY)


def save_report_pdf(company_name: str, markdown_text: str) -> Path:
    out_path = REPORTS_DIR / f"report_{normalize_name(company_name)}.pdf"

    html_text = md.markdown(
        markdown_text,
        extensions=["tables", "fenced_code"]
    )

    css = CSS(string="""
        h1 { font-size: 22pt; font-weight: bold; margin-bottom: 10pt; }
        h2 { font-size: 18pt; font-weight: bold; margin-top: 15pt; margin-bottom: 8pt; }
        h3 { font-size: 14pt; font-weight: bold; margin-top: 12pt; margin-bottom: 6pt; }
        p { font-size: 11pt; line-height: 1.5; }
        ul, ol { margin-left: 20pt; }
        table { border-collapse: collapse; width: 100%; margin: 12pt 0; }
        th, td { border: 1px solid #666; padding: 6pt; font-size: 10pt; }
    """)

    HTML(string=html_text).write_pdf(str(out_path), stylesheets=[css])
    return out_path



def search_online_esg(company_name: str, topics: List[str], max_results: int = 5) -> List[LCDocument]:
    docs = []
    for topic in topics:
        query = f"{company_name} ESG {topic}"
        try:
            res = tavily_client.search(query=query, num_results=max_results)
            for r in res["results"]:
                content = r.get("content", "") or r.get("snippet", "")
                if content.strip():
                    docs.append(
                        LCDocument(
                            page_content=content,
                            metadata={
                                "company": company_name,
                                "source": "tavily",
                                "url": r.get("url", ""),
                                "title": r.get("title", topic)
                            }
                        )
                    )
        except Exception as e:
            print(f"[WARN] Tavily search failed for query: {query}, error: {e}")
    return docs

def extract_company_descriptions(docx_path: Path) -> Dict[str, str]:
    doc = Document(docx_path)
    companies, current_company, buffer = {}, None, []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("Company"):
            if current_company and buffer:
                companies[current_company] = "\n".join(buffer).strip()
                buffer = []
            if ":" in text:
                current_company = text.split(":", 1)[1].strip()
        else:
            buffer.append(text)
    if current_company and buffer:
        companies[current_company] = "\n".join(buffer).strip()
    return companies

def load_first_sheet(xlsx_path: Path) -> pd.DataFrame:
    df_all = pd.read_excel(xlsx_path, sheet_name=None)
    first = list(df_all.keys())[0]
    return df_all[first].copy()

def normalize_name(name: str) -> str:
    return re.sub(r"\s+", "", name.strip().lower())

# def build_documents(company_name: str, description: str, insights_df: pd.DataFrame) -> List[LCDocument]:
#     docs = [LCDocument(page_content=description, metadata={"company": company_name, "source": "docx"})]
#     mask = insights_df["company"].str.lower().str.contains(normalize_name(company_name), na=False) \
#            | insights_df["company"].str.lower().eq(company_name.lower())
#     subset = insights_df[mask].copy() if "company" in insights_df.columns else pd.DataFrame()
#     for i, r in enumerate(subset.itertuples(index=False)):
#         parts = []
#         if hasattr(r, "name") and getattr(r, "name"):
#             parts.append(f"{getattr(r,'name')}")
#         if hasattr(r, "description") and getattr(r, "description"):
#             parts.append(f"{getattr(r,'description')}")
#         if hasattr(r, "context") and isinstance(getattr(r,"context"), str) and getattr(r,"context").strip():
#             parts.append(getattr(r,"context").strip())
#         content = " | ".join(p for p in parts if p)
#         md = {
#             "company": company_name,
#             "source": f"xlsx-{i}",
#             "file_name": getattr(r, "file_name", ""),
#             "page": getattr(r, "page_nbr", ""),
#         }
#         if content:
#             docs.append(LCDocument(page_content=content, metadata=md))
#     return docs
def build_documents(company_name: str, description: str, insights_df: pd.DataFrame) -> List[LCDocument]:
    docs = [LCDocument(page_content=description, metadata={"company": company_name, "source": "docx"})]

    # insights
    if "company" in insights_df.columns:
        mask = insights_df["company"].str.lower().str.contains(normalize_name(company_name), na=False) \
               | insights_df["company"].str.lower().eq(company_name.lower())
        subset = insights_df[mask].copy()
    else:
        subset = pd.DataFrame()

    for i, r in enumerate(subset.itertuples(index=False)):
        parts = []
        if hasattr(r, "name") and getattr(r, "name"):
            parts.append(f"{getattr(r,'name')}")
        if hasattr(r, "description") and getattr(r, "description"):
            parts.append(f"{getattr(r,'description')}")
        if hasattr(r, "context") and isinstance(getattr(r,"context"), str) and getattr(r,"context").strip():
            parts.append(getattr(r,"context").strip())
        content = " | ".join(p for p in parts if p)
        md = {
            "company": company_name,
            "source": f"xlsx-{i}",
            "file_name": getattr(r, "file_name", ""),
            "page": getattr(r, "page_nbr", ""),
        }
        if content:
            docs.append(LCDocument(page_content=content, metadata=md))

    # online ESG data
    esg_topics = ["climate change", "pollution", "water use", "biodiversity", "circular economy",
                  "workforce", "supply chain labor", "communities", "consumers"]
    online_docs = search_online_esg(company_name, esg_topics, max_results=3)
    docs.extend(online_docs)

    return docs

def build_or_load_vs(company_name: str, docs: List[LCDocument], emb: AzureOpenAIEmbeddings) -> FAISS:
    vs_path = VSTORE_DIR / f"vs_{normalize_name(company_name)}"
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500, chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    splits = splitter.split_documents(docs)
    # Lightweight dedup by content hash
    seen = set(); unique = []
    for d in splits:
        h = hash(d.page_content.strip())
        if h not in seen:
            seen.add(h); unique.append(d)

    if vs_path.exists():
        store = FAISS.load_local(vs_path, emb, allow_dangerous_deserialization=True)
    else:
        store = FAISS.from_documents(unique, emb)
        store.save_local(vs_path)
    return store

SECTION_PLANNER_TMPL = PromptTemplate(
    input_variables=["company_name","industry_hint","doc_summary"],
    template=(
        "You are an ESG reporting planner. Propose a stable, company-specific outline for a 4 to 6 page Markdown report, focusing ONLY on environmental and social impacts. o NOT include governance.  "
        "aligned with ESRS. Include a numbered list of sections with titles and 2 to 5 bullet subpoints each. "
        "Emphasize both actual impacts and potential impacts for environmental and social topics. "
        "Also include a table plan and where to place it. Keep it deterministic and avoid reusing identical text.\n\n"
        "Company: {company_name}\n"
        "Industry hint: {industry_hint}\n"
        "Context summary:\n{doc_summary}\n\n"
        "Return JSON with keys: sections, each item has title, goals, tokens_target; tables, each item has title, purpose, fields."
    )
)

def summarize_docs_for_planning(vs: FAISS, company_query: str = "high level ESG context and industry positioning") -> str:
    # Shallow retrieval to create a planning summary
    docs = vs.similarity_search(company_query, k=8)
    text = "\n".join(d.page_content for d in docs)
    return text[:6000]

def plan_sections(company_name: str, industry_hint: str, vs: FAISS) -> Dict:
    summ = summarize_docs_for_planning(vs)
    plan_msg = SECTION_PLANNER_TMPL.format(
        company_name=company_name,
        industry_hint=industry_hint,
        doc_summary=summ
    )
    resp = llm.invoke(plan_msg)
    # Try JSON extraction
    m = re.search(r"\{.*\}", resp.content, flags=re.S)
    plan_json = json.loads(m.group(0)) if m else {"sections": [], "tables": []}
    # Safety defaults
    for s in plan_json.get("sections", []):
        s.setdefault("tokens_target", 900)
        s.setdefault("goals", [])
    return plan_json


QUERY_ANALYZER_TMPL = PromptTemplate(
    input_variables=["section_title","goals"],
    template=(
        "Convert the section intent to 2 retrieval queries and 1 boolean filter suggestion over metadata "
        "(e.g., source contains xlsx or docx, or page range). Keep queries concise.\n"
        "Section: {section_title}\nGoals: {goals}\n"
        "Return JSON with keys: queries, filters"
    )
)

WRITER_TMPL = PromptTemplate(
    input_variables=["company","section","goals","evidence","tables_plan","tokens_target"],
    template=(
        "Write the section as Markdown for an ESG report compliant with ESRS, professional narrative, "
        "2 to 4 paragraphs plus if useful a compact table or bullet list. "
        "Explicitly cover Actual impacts and Potential impacts. "
        "Weave quantitative and qualitative evidence from the provided context only. "
        "Conclude with Gaps, and Recommendations with 2 to 3 items.\n\n"
        "Company: {company}\nSection: {section}\nGoals: {goals}\n"
        "Tables allowable: {tables_plan}\n"
        "Context evidence:\n{evidence}\n\n"
        "Target tokens: {tokens_target}\n"
        "Output only the Markdown for this section."
    )
)

def retrieve_for_section(vs: FAISS, section_title: str, goals: List[str]) -> List[LCDocument]:
    qa = QUERY_ANALYZER_TMPL.format(section_title=section_title, goals="\n".join(goals))
    resp = llm.invoke(qa)
    try:
        spec = json.loads(re.search(r"\{.*\}", resp.content, flags=re.S).group(0))
    except Exception:
        spec = {"queries": [section_title], "filters": {}}

    results = []
    for q in spec.get("queries", [])[:3]:
        results.extend(vs.similarity_search(q, k=6))
    # Simple rerank by unique content and source diversity
    uniq, seen = [], set()
    for d in results:
        key = (hash(d.page_content), d.metadata.get("source",""))
        if key not in seen:
            seen.add(key); uniq.append(d)
    return uniq[:12]

def write_section(company: str, section: Dict, vs: FAISS, tables_plan: List[Dict]) -> str:
    docs = retrieve_for_section(vs, section["title"], section.get("goals", []))
    evidence = "\n\n".join(f"- {d.page_content}" for d in docs)
    tables_ok = [t for t in tables_plan if section["title"].lower() in t.get("purpose","").lower() or section["title"].lower() in t.get("title","").lower()]
    msg = WRITER_TMPL.format(
        company=company,
        section=section["title"],
        goals="\n".join(section.get("goals", [])),
        evidence=evidence[:8000],
        tables_plan=json.dumps(tables_ok[:2]),
        tokens_target=section.get("tokens_target", 900)
    )
    out = llm.invoke(msg).content
    return out, docs

def build_report(company_name: str, vs: FAISS, industry_hint: str = "Unknown") -> str:
    plan = plan_sections(company_name, industry_hint, vs)
    sections = plan.get("sections", [])
    tables_plan = plan.get("tables", [])

    md = [f"# ESG Report, {company_name}", "", "## Contents"]
    for i, s in enumerate(sections, 1):
        anchor = re.sub(r'[^a-z0-9]+', '-', s["title"].strip().lower())
        md.append(f"{i}. [{s['title']}](#{anchor})")
    md.append("")

    all_refs = []
    for s in sections:
        # md.append(f"## {s['title']}")
        body, used_docs = write_section(company_name, s, vs, tables_plan)
        md.append(body.strip())
        # refs
        # refs = []
        # for idx, d in enumerate(used_docs, 1):
        #     refs.append(f"[{idx}] {d.metadata.get('file_name','')}, p.{d.metadata.get('page','')}, {d.metadata.get('source','')}")
        refs = []
        for idx, d in enumerate(used_docs, 1):
            if d.metadata.get("source") == "tavily":
                refs.append(
                    f"[{idx}] {d.metadata.get('title', 'online source')} ({d.metadata.get('url', '')}) — Tavily")
            else:
                refs.append(
                    f"[{idx}] {d.metadata.get('file_name', '')}, p.{d.metadata.get('page', '')}, {d.metadata.get('source', '')}")

        if refs:
            md.append("")
            md.append("### References")
            md.extend(refs)
            all_refs.extend(refs)
        md.append("")

    md.append("## Appendix, Methodology and Limitations")
    md.append("This report uses retrieval augmented generation, context limited to provided sources, "
              "length control per section, and ESRS framing for actual and potential impacts. "
              "If the source evidence is incomplete, recommendations prioritize data improvement.")

    return "\n".join(md)

def save_report(company_name: str, markdown: str) -> Path:
    out_path = REPORTS_DIR / f"report_{normalize_name(company_name)}.md"
    out_path.write_text(markdown, encoding="utf-8")
    return out_path

if __name__ == '__main__':
    # Detect input files
    DOCX_PATH = next(DATA_DIR.glob("TO_SHARE_ai_engineering_challenge_company_descriptions*.docx"))
    XLSX_PATH = next(DATA_DIR.glob("TO_SHARE_insights_extract_ai_engineering_challenge*.xlsx"))

    company_map = extract_company_descriptions(DOCX_PATH)
    insights_df = load_first_sheet(XLSX_PATH)

    for company_name, description in list(company_map.items()):
        print(f"\n=== Building ESG report for {company_name} ===")
        docs = build_documents(company_name, description, insights_df)
        vs = build_or_load_vs(company_name, docs, embeddings)
        report_md = build_report(company_name, vs, industry_hint="Logistics and Supply Chain")

        # Save Markdown
        path_md = save_report(company_name, report_md)
        print(f"Saved Markdown to {path_md}")

        # Save PDF
        path_pdf = save_report_pdf(company_name, report_md)
        print(f"Saved PDF to {path_pdf}")
